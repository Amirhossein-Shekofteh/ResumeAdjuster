from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader

from src.config import CONFIG


class DocumentLoadingError(Exception):
    """
    Raised when a document cannot be loaded or has an unsupported format.
    """


DocumentSource = str | Path | BinaryIO


def _is_path_like(source: DocumentSource) -> bool:
    return isinstance(source, str | Path)


def _get_extension(source: DocumentSource, filename: str | None = None) -> str:
    """
    Get the file extension from a path, explicit filename, or file-like object.

    Streamlit uploaded files usually have a `.name` attribute, but this function
    also accepts an explicit filename for reliability.
    """

    if filename:
        extension = Path(filename).suffix.lower()
    elif _is_path_like(source):
        extension = Path(source).suffix.lower()
    else:
        source_name = getattr(source, "name", None)
        extension = Path(source_name).suffix.lower() if source_name else ""

    if not extension:
        raise DocumentLoadingError(
            "Could not determine file type. Provide a filename with .txt, .pdf, or .docx."
        )

    return extension


def _validate_extension(extension: str) -> None:
    """
    Validate that the extension is allowed by project configuration.
    """

    if extension not in CONFIG.allowed_file_types:
        allowed = ", ".join(CONFIG.allowed_file_types)
        raise DocumentLoadingError(
            f"Unsupported file type: {extension}. Allowed file types: {allowed}."
        )


def _read_binary(source: DocumentSource) -> bytes:
    """
    Read bytes from a path or file-like object.
    """

    if _is_path_like(source):
        return Path(source).read_bytes()

    try:
        source.seek(0)
    except (AttributeError, OSError):
        pass

    data = source.read()

    if isinstance(data, str):
        return data.encode("utf-8")

    if isinstance(data, bytes):
        return data

    raise DocumentLoadingError(
        f"Could not read binary data from source. Got {type(data).__name__}."
    )


def _decode_text(data: bytes) -> str:
    """
    Decode text bytes using common encodings.

    Most text files should be UTF-8, but this fallback helps with copied or
    exported resumes that may use another encoding.
    """

    encodings = ("utf-8", "utf-8-sig", "utf-16", "latin-1")

    for encoding in encodings:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    return data.decode("utf-8", errors="replace")


def load_txt(source: DocumentSource) -> str:
    """
    Load a plain text file.

    Supports:
    - Local path
    - Streamlit uploaded file
    - Any binary file-like object
    """

    data = _read_binary(source)
    text = _decode_text(data)

    return text.strip()


def load_pdf(source: DocumentSource) -> str:
    """
    Load text from a PDF file.

    This function extracts embedded text from PDFs. It does not perform OCR, so
    scanned image-only PDFs may return little or no text.
    """

    try:
        if _is_path_like(source):
            reader = PdfReader(str(source))
        else:
            data = _read_binary(source)
            reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise DocumentLoadingError(f"Could not open PDF file: {exc}") from exc

    page_texts: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            raise DocumentLoadingError(
                f"Could not extract text from PDF page {index}: {exc}"
            ) from exc

        text = text.strip()

        if text:
            page_texts.append(text)

    extracted_text = "\n\n".join(page_texts).strip()

    if not extracted_text:
        raise DocumentLoadingError(
            "No readable text was found in the PDF. "
            "The PDF may be scanned or image-based. Try using a text-based PDF, DOCX, or TXT file."
        )

    return extracted_text


def _extract_docx_table_text(document: Document) -> list[str]:
    """
    Extract readable text from DOCX tables.

    Resumes often use tables for layout, so ignoring tables can lose important
    content.
    """

    table_texts: list[str] = []

    for table in document.tables:
        for row in table.rows:
            cell_values: list[str] = []

            for cell in row.cells:
                cell_text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )

                if cell_text:
                    cell_values.append(cell_text)

            if cell_values:
                table_texts.append(" | ".join(cell_values))

    return table_texts


def load_docx(source: DocumentSource) -> str:
    """
    Load text from a DOCX file.

    Extracts both normal paragraphs and table content.
    """

    try:
        if _is_path_like(source):
            document = Document(str(source))
        else:
            data = _read_binary(source)
            document = Document(BytesIO(data))
    except Exception as exc:
        raise DocumentLoadingError(f"Could not open DOCX file: {exc}") from exc

    paragraph_texts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_texts = _extract_docx_table_text(document)

    extracted_text = "\n".join(paragraph_texts + table_texts).strip()

    if not extracted_text:
        raise DocumentLoadingError("No readable text was found in the DOCX file.")

    return extracted_text


def load_document(source: DocumentSource, filename: str | None = None) -> str:
    """
    Load text from TXT, PDF, or DOCX.

    For local files:

        load_document("data/sample_resume.txt")

    For Streamlit uploaded files:

        load_document(uploaded_file, filename=uploaded_file.name)
    """

    extension = _get_extension(source, filename=filename)
    _validate_extension(extension)

    if extension == ".txt":
        return load_txt(source)

    if extension == ".pdf":
        return load_pdf(source)

    if extension == ".docx":
        return load_docx(source)

    raise DocumentLoadingError(f"No loader implemented for file type: {extension}.")